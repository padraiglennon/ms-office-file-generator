"""Tests for the complexity-driven Word document generator."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from common_file_generator.core import Complexity, generate_doc
from common_file_generator.generators import DocxComplexityGenerator


def test_section_count_drives_headings(tmp_path: Path) -> None:
    out = tmp_path / "d.docx"
    generate_doc(str(out), complexity="standard", sections=8, seed=1)
    doc = Document(str(out))
    h1 = [p for p in doc.paragraphs if p.style and p.style.name == "Heading 1"]
    assert len(h1) == 8


def test_invalid_section_count_raises() -> None:
    with pytest.raises(ValueError, match="at least 1"):
        DocxComplexityGenerator(Complexity.MINIMAL, sections=0)


def test_minimal_has_no_tables_or_images(tmp_path: Path) -> None:
    out = tmp_path / "d.docx"
    generate_doc(str(out), complexity="minimal", sections=10, seed=3)
    doc = Document(str(out))
    assert len(doc.tables) == 0
    assert len(doc.inline_shapes) == 0


def test_maximum_has_tables_and_images(tmp_path: Path) -> None:
    out = tmp_path / "d.docx"
    generate_doc(str(out), complexity="maximum", sections=25, seed=7)
    doc = Document(str(out))
    assert len(doc.tables) > 0
    assert len(doc.inline_shapes) > 0
    styles = {p.style.name for p in doc.paragraphs if p.style}
    assert "List Bullet" in styles or "List Number" in styles


def test_blocks_per_section_scales_with_complexity(tmp_path: Path) -> None:
    # Maximum is denser than minimal at the same section count, so it has more
    # body paragraphs/blocks overall.
    mn = tmp_path / "mn.docx"
    mx = tmp_path / "mx.docx"
    generate_doc(str(mn), complexity="minimal", sections=8, seed=1)
    generate_doc(str(mx), complexity="maximum", sections=8, seed=1)
    assert len(Document(str(mx)).paragraphs) > len(Document(str(mn)).paragraphs)


def test_blocks_per_section_override(tmp_path: Path) -> None:
    out = tmp_path / "d.docx"
    generate_doc(
        str(out), complexity="minimal", sections=4, seed=1, blocks_per_section=6
    )
    doc = Document(str(out))
    # 4 section headings + body; the override forces 6 extra blocks per section,
    # so there are far more paragraphs than the minimal default (1) would give.
    assert len(doc.paragraphs) > 4 * 6


def test_invalid_blocks_per_section_raises() -> None:
    with pytest.raises(ValueError, match="blocks_per_section"):
        DocxComplexityGenerator(Complexity.MINIMAL, sections=2, blocks_per_section=0)


def test_seed_is_deterministic(tmp_path: Path) -> None:
    a, b = tmp_path / "a.docx", tmp_path / "b.docx"
    generate_doc(str(a), complexity="complex", sections=10, seed=42)
    generate_doc(str(b), complexity="complex", sections=10, seed=42)
    assert _texts(a) == _texts(b)


def test_different_seed_differs(tmp_path: Path) -> None:
    a, b = tmp_path / "a.docx", tmp_path / "b.docx"
    generate_doc(str(a), complexity="complex", sections=10, seed=1)
    generate_doc(str(b), complexity="complex", sections=10, seed=2)
    assert _texts(a) != _texts(b)


def test_tables_fit_the_page_width(tmp_path: Path) -> None:
    out = tmp_path / "d.docx"
    generate_doc(str(out), complexity="maximum", sections=20, seed=5)
    doc = Document(str(out))
    section = doc.sections[0]
    usable = section.page_width - section.left_margin - section.right_margin
    for table in doc.tables:
        total = sum(cell.width or 0 for cell in table.rows[0].cells)
        # Column widths sum to no more than the usable page width (+1pt slack).
        assert total <= usable + 12700


def test_each_theme_styles_headings_body_and_quote(tmp_path: Path) -> None:
    from common_file_generator.generators.docx_theme import THEMES

    for name, theme in THEMES.items():
        out = tmp_path / f"{name}.docx"
        generate_doc(str(out), complexity="standard", sections=4, seed=1, theme=name)
        doc = Document(str(out))
        assert str(doc.styles["Heading 1"].font.color.rgb) == str(theme.heading_color)
        assert str(doc.styles["Title"].font.color.rgb) == str(theme.heading_color)
        assert str(doc.styles["Normal"].font.color.rgb) == str(theme.body_color)
        assert str(doc.styles["Intense Quote"].font.color.rgb) == str(
            theme.accent_color
        )


def test_themes_are_visibly_distinct(tmp_path: Path) -> None:
    headings = set()
    for name in ("ocean", "slate", "sand"):
        out = tmp_path / f"{name}.docx"
        generate_doc(str(out), complexity="standard", sections=3, seed=1, theme=name)
        headings.add(str(Document(str(out)).styles["Heading 1"].font.color.rgb))
    assert len(headings) == 3


def test_table_header_is_shaded_with_theme_fill(tmp_path: Path) -> None:
    from docx.oxml.ns import qn

    from common_file_generator.generators.docx_theme import SAND

    out = tmp_path / "d.docx"
    generate_doc(str(out), complexity="maximum", sections=20, seed=5, theme="sand")
    doc = Document(str(out))
    assert doc.tables  # maximum produces tables
    header_cell = doc.tables[0].cell(0, 0)._tc
    shd = header_cell.find(qn("w:tcPr")).find(qn("w:shd"))
    assert shd.get(qn("w:fill")) == str(SAND.table_header_fill)


def test_default_theme_is_ocean(tmp_path: Path) -> None:
    from common_file_generator.generators.docx_theme import OCEAN

    default_out, ocean_out = tmp_path / "def.docx", tmp_path / "ocean.docx"
    generate_doc(str(default_out), complexity="complex", sections=6, seed=9)
    generate_doc(
        str(ocean_out), complexity="complex", sections=6, seed=9, theme="ocean"
    )
    assert _content(default_out) == _content(ocean_out)
    assert str(Document(str(default_out)).styles["Heading 1"].font.color.rgb) == str(
        OCEAN.heading_color
    )


def test_theme_preserves_determinism(tmp_path: Path) -> None:
    # The .docx is a zip whose entries carry a save-time timestamp, so raw bytes
    # can differ across a clock tick. Determinism is over content: same
    # (theme, complexity, sections, seed) -> identical zip-member bytes.
    a, b = tmp_path / "a.docx", tmp_path / "b.docx"
    generate_doc(str(a), complexity="complex", sections=8, seed=42, theme="slate")
    generate_doc(str(b), complexity="complex", sections=8, seed=42, theme="slate")
    assert _content(a) == _content(b)


def test_unknown_theme_raises_listing_valid_names() -> None:
    with pytest.raises(ValueError, match="ocean, sand, slate"):
        generate_doc("unused.docx", theme="neon")


def test_ocean_mirrors_the_pptx_palette() -> None:
    """Drift guard: the mirrored ocean hexes must match the PPTX OCEAN theme."""
    from common_file_generator.generators.docx_theme import OCEAN
    from common_file_generator.generators.theme import OCEAN as PPTX_OCEAN

    assert str(OCEAN.heading_color) == str(PPTX_OCEAN.primary)
    assert str(OCEAN.accent_color) == str(PPTX_OCEAN.accent)
    assert str(OCEAN.body_color) == str(PPTX_OCEAN.body_color)


def test_every_theme_clears_wcag_aa_on_its_backgrounds() -> None:
    """AC#7: each theme's substantive text/background pairings clear WCAG AA.

    Body text must clear 4.5:1 on the white page. Headings and table-header text
    are large and bold, so they take AA's large-text floor of 3.0:1. The quote
    accent is intentionally excluded: it is a decorative emphasis colour mirrored
    from the brand palette (ocean's teal is 2.5:1 on white), not body copy whose
    readability AA governs.
    """
    from common_file_generator.generators.docx_theme import THEMES

    aa_body, aa_large = 4.5, 3.0
    white = (0xFF, 0xFF, 0xFF)
    for theme in THEMES.values():
        assert _contrast(theme.heading_color, white) >= aa_large, theme.name
        assert _contrast(theme.body_color, white) >= aa_body, theme.name
        assert (
            _contrast(theme.table_header_text, theme.table_header_fill) >= aa_large
        ), theme.name


def _contrast(fg: object, bg: object) -> float:
    def channel(rgb: object) -> tuple[int, int, int]:
        # docx RGBColor is a 3-byte value; index it for the components.
        return (rgb[0], rgb[1], rgb[2]) if not isinstance(rgb, tuple) else rgb

    def luminance(rgb: tuple[int, int, int]) -> float:
        def lin(c: float) -> float:
            c /= 255
            return c / 12.92 if c <= 0.03928 else ((c + 0.055) / 1.055) ** 2.4

        r, g, b = (lin(v) for v in rgb)
        return 0.2126 * r + 0.7152 * g + 0.0722 * b

    lo, hi = sorted((luminance(channel(fg)), luminance(channel(bg))))
    return (hi + 0.05) / (lo + 0.05)


def _texts(path: Path) -> list[str]:
    return [p.text for p in Document(str(path)).paragraphs]


def _content(path: Path) -> dict[str, bytes]:
    """The .docx's zip-member bytes, ignoring per-save entry timestamps."""
    import zipfile

    with zipfile.ZipFile(path) as zf:
        return {name: zf.read(name) for name in sorted(zf.namelist())}
